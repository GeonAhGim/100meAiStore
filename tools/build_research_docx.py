from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "report-source.md"
OUTPUT = ROOT / "docs" / "AI_위탁판매_자동운영_아키텍처_사전조사.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9D9D9", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "245B78")
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, text):
    url_pattern = re.compile(r"https?://[^\s;]+")
    pos = 0
    for match in url_pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        url = match.group(0).rstrip(".,)")
        add_hyperlink(paragraph, url, url)
        pos = match.start() + len(match.group(0))
    if pos < len(text):
        paragraph.add_run(text[pos:])


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def keep_table_row_together(row, repeat_header=False):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    if repeat_header:
        tr_pr.append(OxmlElement("w:tblHeader"))


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.13
    for name, size, before, after in (
        ("Title", 24, 0, 16), ("Heading 1", 16, 14, 6), ("Heading 2", 12.5, 10, 4),
        ("Heading 3", 11, 7, 3)
    ):
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    title_style_ppr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_style_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_style_ppr.remove(title_border)

    i = 0
    title_done = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# ") and not title_done:
            p = doc.add_paragraph(line[2:], style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_done = True
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            i += 1
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(7)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8.2)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
            table_lines = [line]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
            is_ledger = len(rows[0]) >= 6
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(rows):
                keep_table_row_together(table.rows[r_idx], repeat_header=(r_idx == 0))
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_border(cell)
                    set_cell_margins(cell, top=55, start=70, bottom=55, end=70) if is_ledger else set_cell_margins(cell)
                    if r_idx == 0:
                        set_cell_shading(cell, "203B53")
                    elif r_idx % 2 == 0:
                        set_cell_shading(cell, "EDF3F7")
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0 if is_ledger else 1.08
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.add_run(value)
                    run.bold = r_idx == 0
                    run.font.size = Pt(7.2 if is_ledger else (8.4 if len(rows[0]) >= 5 else 9))
                    run.font.color.rgb = RGBColor(255, 255, 255) if r_idx == 0 else RGBColor(0, 0, 0)
            doc.add_paragraph().paragraph_format.space_after = Pt(1)
            continue
        numbered = re.match(r"^(\d+)\. (.*)", line)
        if numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.22)
            p.add_run(f"{numbered.group(1)}. ")
            add_rich_text(p, numbered.group(2))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        add_rich_text(p, line.replace("  ", " "))
        i += 1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AI 위탁판매 자동운영 아키텍처 사전조사  |  2026-09-04")
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.core_properties.title = "AI 위탁판매 자동운영 아키텍처 사전조사"
    doc.core_properties.subject = "OpenAI Codex와 ChatGPT를 운영 에이전트로 통합하기 위한 리테일 및 오픈소스 조사"
    doc.core_properties.author = ""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
